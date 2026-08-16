"""Assembles reports/analysis_report.{md,html,docx} from figures.py + tables.py.

`build_report()` is the single entry point scripts/run_report.py calls. It
runs every figure/table function (cheap: they read cached artefacts, they do
not re-run training or feature extraction -- see the module docstrings in
figures.py/tables.py), then renders three parallel views of the SAME content
so nothing drifts between formats:

    reports/analysis_report.md    plain markdown, images as relative links
    reports/analysis_report.html  self-contained, images inlined as base64
    reports/analysis_report.docx  navy/steel blue, Calibri, justified body

The report's content -- every paragraph of prose -- lives in `build_sections`
below as a list of `Section` objects. All three renderers just walk that
list, so editing report text means editing one place.
"""

from __future__ import annotations

import base64
import hashlib
import importlib.metadata as im
import json
import subprocess
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd

from . import figures as figs
from . import tables as tbls
from .style import NAVY, STEEL_BLUE

# Every third-party library the pipeline actually uses, with the role it plays.
# Listed exhaustively so the software environment is reproducible from the
# report alone, without reading the source.
SOFTWARE: list[tuple[str, str, str]] = [
    # (distribution, import name, role)
    ("numpy", "numpy", "Array computation throughout; all coordinate and "
                       "embedding maths"),
    ("pandas", "pandas", "Tabular data: single-cell tables, patch labels, "
                         "feature matrices"),
    ("scipy", "scipy", "cKDTree for 40 um neighbourhood and nearest-neighbour "
                       "queries; sparse matrices; Wilcoxon signed-rank test; "
                       "FFT cross-correlation for registration QC"),
    ("scikit-learn", "sklearn", "k-means, latent Dirichlet allocation, "
                                "silhouette / Davies-Bouldin / adjusted Rand "
                                "index, agglomerative consensus clustering, "
                                "precision/recall/F1 and confusion matrices"),
    ("pyyaml", "yaml", "Configuration files and the inherit-chain resolver"),
    ("pyarrow", "pyarrow", "Parquet engine for every cached artefact"),
    ("tqdm", "tqdm", "Progress reporting in long stage loops"),
    ("tabulate", "tabulate", "Markdown table rendering in the report"),

    ("lifelines", "lifelines", "Cox proportional hazards models, Kaplan-Meier "
                               "estimation, log-rank tests, concordance index"),
    ("scikit-survival", "sksurv", "LASSO-Cox (CoxnetSurvivalAnalysis), random "
                                  "survival forest permutation importance, "
                                  "time-dependent cumulative AUC"),
    ("statsmodels", "statsmodels", "Benjamini-Hochberg false discovery rate "
                                   "correction"),

    ("networkx", "networkx", "Connected-component fallback for collinearity "
                             "reduction"),
    ("python-igraph", "igraph", "Louvain community detection on the feature "
                                "correlation graph"),

    ("torch", "torch", "Neural network components, where used: the classifier head, focal "
                       "loss, weighted sampling, and the graph/2D-grid/3D-grid "
                       "context encoders"),
    ("torchvision", "torchvision", "ResNet-50 benchmark encoder and its own "
                                   "preprocessing transforms"),
    ("timm", "timm", "Model registry and data configuration for UNI and MUSK"),
    ("transformers", "transformers", "Phikon encoder (ViTModel) and its "
                                     "AutoImageProcessor"),
    ("huggingface-hub", "huggingface_hub", "Weight download and token-gated "
                                           "model access"),

    ("pillow", "PIL", "Patch image handling and resampling"),
    ("opencv-python-headless", "cv2", "Tissue segmentation (HSV, Otsu, "
                                      "morphology), artefact filtering, "
                                      "marker gating thresholds"),
    ("tifffile", "tifffile", "Pyramidal OME-TIFF reading for the paired H&E "
                             "whole-slide images"),
    ("imagecodecs", "imagecodecs", "JPEG and zlib codecs those OME-TIFFs use"),
    ("zarr", "zarr", "Windowed level-0 reads, so a 13.5 GB image is never "
                     "loaded whole"),

    ("matplotlib", "matplotlib", "Every figure in this report"),
    ("python-docx", "docx", "Word rendering of the report and manuscript"),
    ("openpyxl", "openpyxl", "Reads the TCGA-CDR clinical workbook"),

    ("pytest", "pytest", "Test suite"),
]

PACKAGES = [dist for dist, _imp, _role in SOFTWARE]

# Deliberately optional. Each has a documented fallback so the pipeline runs
# without it, and the fallback taken is recorded.
OPTIONAL_SOFTWARE: list[tuple[str, str]] = [
    ("spatial-lda", "True spatial-LDA prior with a smoothness penalty over "
                    "neighbouring index cells. NOT installed in this run, so "
                    "scikit-learn's LatentDirichletAllocation was used instead: "
                    "the same generative model without the spatial prior. This "
                    "is a real deviation and is recorded in T10."),
    ("openslide-python", "Aperio/SVS reading for TCGA slides. Not required for "
                         "the Orion cohort, whose OME-TIFFs are read with "
                         "tifffile and zarr."),
    ("stardist", "Nuclear segmentation. Not required here: the paired cohort "
                 "ships its own segmented single-cell tables."),
]




# ============================================================== run everything


def _run_all(figures_dir: str, tables_dir: str) -> tuple[dict, dict]:
    fig = {m["id"]: m for m in (fn(figures_dir) for fn in figs.ALL_FIGURES)}
    tab = {m["id"]: m for m in (fn(tables_dir) for fn in tbls.ALL_TABLES)}
    return fig, tab


def _reproducibility_stats(config_path: str = "config/coda_params.yaml") -> dict:
    """Environment and provenance. The parameter hash comes from the project's
    own guard, so the value reported here is the one that gates every run."""
    import platform
    import sys as _sys

    root = Path(__file__).resolve().parents[3]
    _sys.path.insert(0, str(root / "src"))
    from coda_my.guard import flatten, verify

    params = verify()                       # raises if any locked value drifted
    n_locked = len(flatten(params["locked"]))
    cfg_hash = hashlib.sha256(
        json.dumps(params["locked"], sort_keys=True, default=str).encode()).hexdigest()

    versions = {}
    for pkg in PACKAGES:
        try:
            versions[pkg] = im.version(pkg)
        except im.PackageNotFoundError:
            versions[pkg] = "not installed"

    try:
        sha = subprocess.run(["git", "rev-parse", "HEAD"], capture_output=True,
                             text=True, timeout=5, cwd=str(root)).stdout.strip()
        git_sha = sha if sha else "NOT A GIT REPOSITORY"
    except Exception:
        git_sha = "NOT A GIT REPOSITORY"

    return {"config_path": config_path, "config_hash_sha256": cfg_hash[:16],
            "n_locked_parameters": n_locked,
            "versions": versions, "git_sha": git_sha,
            "python_version": platform.python_version(),
            "platform": f"{platform.system()} {platform.release()}",
            "project_seed": params.get("locked", {}).get("statistics", {})
                                  .get("seed", 42)}


# ==================================================================== content


from ._sections import Section, build_sections  # noqa: E402,F401


def _software_table(versions: dict[str, str]) -> str:
    """Markdown table of every library, its version and what it is used for."""
    rows = ["| library | version | import | role |", "|---|---|---|---|"]
    for dist, imp, role in SOFTWARE:
        rows.append(f"| `{dist}` | {versions.get(dist, 'not installed')} | "
                    f"`{imp}` | {role} |")
    return "\n".join(rows)


# ==================================================================== markdown


def _df_to_md(df: pd.DataFrame, max_rows: int = 25) -> str:
    if len(df) > max_rows:
        shown = df.head(max_rows)
        note = f"\n\n*(showing {max_rows} of {len(df)} rows; full table in the CSV)*"
    else:
        shown, note = df, ""
    return shown.to_markdown(index=False, floatfmt=".4g") + note


def _inject_software(sections, versions):
    """Replace the placeholder paragraph with the rendered software table."""
    for sec in sections:
        sec.paragraphs = [_software_table(versions) if p == "SOFTWARE_TABLE_PLACEHOLDER"
                          else p for p in sec.paragraphs]
    return sections


def render_markdown(sections: list[Section], fig: dict, tab: dict, figures_dir: str) -> str:
    lines = ["# CODA reproduction: 3D histology and breast IHC histomorphometry", ""]
    for sec in sections:
        lines.append(f"{'#' * (sec.level + 1)} {sec.heading}")
        lines.append("")
        for p in sec.paragraphs:
            lines.append(p)
            lines.append("")
        for fid in sec.figure_ids:
            f = fig[fid]
            rel = Path(f["paths"]["png"]).name
            lines.append(f"**{f['id']}. {f['title']}** -- *{f['source']}*")
            lines.append("")
            lines.append(f"![{f['id']}]({figures_dir}/{rel})")
            lines.append("")
            lines.append(f["caption"])
            lines.append("")
        for tid in sec.table_ids:
            t = tab[tid]
            lines.append(f"**{t['id']}. {t['title']}** -- *{t['source']}*")
            lines.append("")
            lines.append(_df_to_md(t["df"]))
            lines.append("")
            lines.append(t["caption"])
            lines.append("")
    return "\n".join(lines)


# ======================================================================= html


_HTML_HEAD = f"""<!doctype html><html><head><meta charset="utf-8">
<title>CODA reproduction: 3D histology and breast IHC histomorphometry</title>
<style>
body {{ font-family: Calibri, 'Segoe UI', sans-serif; max-width: 960px; margin: 2rem auto;
       padding: 0 1.5rem; line-height: 1.55; color: #1a1a1a; }}
h1 {{ color: {NAVY}; border-bottom: 3px solid {NAVY}; padding-bottom: .3rem; }}
h2 {{ color: {NAVY}; border-bottom: 1px solid #ccc; padding-bottom: .2rem; margin-top: 2.2rem; }}
h3 {{ color: {STEEL_BLUE}; margin-top: 1.6rem; }}
.source-real {{ color: #006400; font-weight: 600; }}
.source-sim {{ color: {STEEL_BLUE}; font-weight: 600; }}
.source-missing {{ color: #b00020; font-weight: 600; }}
.caption {{ font-size: 0.92em; color: #444; margin: .3rem 0 1.2rem 0; }}
img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 4px; }}
table {{ border-collapse: collapse; margin: .8rem 0; font-size: .88em; max-width: 100%; display: block;
        overflow-x: auto; }}
th, td {{ border: 1px solid #ccc; padding: 4px 8px; text-align: left; }}
th {{ background: {NAVY}; color: white; }}
tr:nth-child(even) {{ background: #f4f7fa; }}
code, pre {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
pre {{ padding: 10px; overflow-x: auto; }}
</style></head><body>
<h1>CODA reproduction: 3D histology and breast IHC histomorphometry</h1>
"""


def _md_inline(text: str) -> str:
    """Minimal inline markdown: **bold** and `code`."""
    import re
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    return re.sub(r"`(.+?)`", r"<code>\1</code>", text)


def _md_table_to_html(md: str) -> str:
    """Render a pipe table written in a prose paragraph as a real HTML table."""
    lines = [ln.strip() for ln in md.strip().splitlines() if ln.strip()]
    rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines]
    if len(rows) < 2:
        return f"<p>{md}</p>"
    header, body = rows[0], rows[2:]          # rows[1] is the --- separator
    out = ["<table><thead><tr>"]
    out += [f"<th>{_md_inline(c)}</th>" for c in header]
    out.append("</tr></thead><tbody>")
    for r in body:
        out.append("<tr>" + "".join(f"<td>{_md_inline(c)}</td>" for c in r) + "</tr>")
    out.append("</tbody></table>")
    return "".join(out)


def _md_bullets_to_html(md: str) -> str:
    lead, items = [], []
    for ln in md.splitlines():
        (items if ln.lstrip().startswith("- ") else lead).append(ln)
    html = f"<p>{_md_inline(' '.join(x for x in lead if x.strip()))}</p>" if any(
        x.strip() for x in lead) else ""
    html += "<ul>" + "".join(
        f"<li>{_md_inline(i.lstrip()[2:])}</li>" for i in items) + "</ul>"
    return html


def _source_class(source: str) -> str:
    if source.startswith("REAL"):
        return "source-real"
    if source.startswith("SIMULATED"):
        return "source-sim"
    if source.startswith("MISSING"):
        return "source-missing"
    return ""


def _b64_img(png_path: str) -> str:
    data = base64.b64encode(Path(png_path).read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


def render_html(sections: list[Section], fig: dict, tab: dict) -> str:
    parts = [_HTML_HEAD]
    for sec in sections:
        tag = "h2" if sec.level == 1 else "h3"
        parts.append(f"<{tag}>{sec.heading}</{tag}>")
        for p in sec.paragraphs:
            if p.startswith("```"):
                parts.append(f"<pre>{p.strip('`').lstrip(chr(10))}</pre>")
            elif p.lstrip().startswith("| "):
                parts.append(_md_table_to_html(p))
            elif p.lstrip().startswith("- ") or "\n- " in p:
                parts.append(_md_bullets_to_html(p))
            else:
                parts.append(f"<p>{p}</p>")
        for fid in sec.figure_ids:
            f = fig[fid]
            parts.append(f"<h4>{f['id']}. {f['title']} "
                        f"<span class='{_source_class(f['source'])}'>[{f['source']}]</span></h4>")
            parts.append(f"<img src='{_b64_img(f['paths']['png'])}' alt='{f['id']}'>")
            parts.append(f"<div class='caption'>{f['caption']}</div>")
        for tid in sec.table_ids:
            t = tab[tid]
            parts.append(f"<h4>{t['id']}. {t['title']} "
                        f"<span class='{_source_class(t['source'])}'>[{t['source']}]</span></h4>")
            shown = t["df"].head(25)
            parts.append(shown.to_html(index=False, float_format=lambda x: f"{x:.4g}"))
            if len(t["df"]) > 25:
                parts.append(f"<div class='caption'>(showing 25 of {len(t['df'])} rows; "
                            f"full table at {t['csv_path']})</div>")
            parts.append(f"<div class='caption'>{t['caption']}</div>")
    parts.append("</body></html>")
    return "\n".join(parts)


# ======================================================================= docx


def render_docx(sections: list[Section], fig: dict, tab: dict, out_path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    navy = RGBColor(0x1C, 0x2B, 0x4A)
    steel = RGBColor(0x24, 0x71, 0xA3)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("CODA reproduction: 3D histology and breast IHC histomorphometry", level=0)
    for run in title.runs:
        run.font.color.rgb = navy

    for sec in sections:
        h = doc.add_heading(sec.heading, level=min(sec.level + 1, 4))
        for run in h.runs:
            run.font.color.rgb = navy if sec.level == 1 else steel
        for p_text in sec.paragraphs:
            if p_text.startswith("```"):
                code = p_text.strip("`").lstrip("\n")
                p = doc.add_paragraph(code)
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                continue
            if p_text.lstrip().startswith("| "):
                # a pipe table written in prose becomes a real Word table
                lines = [ln.strip() for ln in p_text.strip().splitlines() if ln.strip()]
                rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines]
                if len(rows) >= 2:
                    header, body = rows[0], rows[2:]
                    tbl = doc.add_table(rows=1, cols=len(header))
                    tbl.style = "Light Grid Accent 1"
                    for i, c in enumerate(header):
                        tbl.rows[0].cells[i].text = c.replace("`", "")
                    for r in body:
                        cells = tbl.add_row().cells
                        for i, c in enumerate(r[:len(header)]):
                            cells[i].text = c.replace("`", "").replace("**", "")
                    doc.add_paragraph()
                    continue
            p = doc.add_paragraph(p_text.replace("**", ""))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for fid in sec.figure_ids:
            f = fig[fid]
            h4 = doc.add_heading(f"{f['id']}. {f['title']}  [{f['source']}]", level=4)
            for run in h4.runs:
                run.font.color.rgb = steel
            doc.add_picture(f["paths"]["png"], width=Cm(15))
            cap = doc.add_paragraph(f["caption"])
            cap.italic = True
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)

        for tid in sec.table_ids:
            t = tab[tid]
            h4 = doc.add_heading(f"{t['id']}. {t['title']}  [{t['source']}]", level=4)
            for run in h4.runs:
                run.font.color.rgb = steel
            shown = t["df"].head(20)
            table = doc.add_table(rows=1, cols=len(shown.columns))
            table.style = "Light Grid Accent 1"
            for i, col in enumerate(shown.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in shown.iterrows():
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val:.4g}" if isinstance(val, float) else str(val)
            if len(t["df"]) > 20:
                doc.add_paragraph(f"(showing 20 of {len(t['df'])} rows; full table at "
                                 f"{t['csv_path']})").italic = True
            cap = doc.add_paragraph(t["caption"])
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ==================================================================== entrypoint


def build_report(figures_dir: str = "figures", tables_dir: str = "results/tables",
                 reports_dir: str = "reports",
                 config_path: str = "config/coda_params.yaml") -> dict:
    fig, tab = _run_all(figures_dir, tables_dir)
    stats = _reproducibility_stats(config_path)
    sections = _inject_software(build_sections(fig, tab, stats), stats["versions"])

    # Completeness contract: EVERY generated figure and table must actually be
    # placed in a section, or it silently never reaches the report. Fail loudly
    # rather than quietly shipping an incomplete document.
    placed_figs = {fid for s in sections for fid in s.figure_ids}
    placed_tabs = {tid for s in sections for tid in s.table_ids}
    orphan_figs = sorted(set(fig) - placed_figs, key=lambda s: int(s[1:]))
    orphan_tabs = sorted(set(tab) - placed_tabs, key=lambda s: int(s[1:]))
    if orphan_figs or orphan_tabs:
        raise AssertionError(
            f"report would omit generated content: figures {orphan_figs}, "
            f"tables {orphan_tabs}. Add them to a Section in build_sections()."
        )

    Path(reports_dir).mkdir(parents=True, exist_ok=True)
    md = render_markdown(sections, fig, tab, figures_dir)
    (Path(reports_dir) / "analysis_report.md").write_text(md, encoding="utf-8")

    html = render_html(sections, fig, tab)
    (Path(reports_dir) / "analysis_report.html").write_text(html, encoding="utf-8")

    render_docx(sections, fig, tab, str(Path(reports_dir) / "analysis_report.docx"))

    n_missing_fig = sum(1 for f in fig.values() if f["source"] == "MISSING DATA")
    n_missing_tab = sum(1 for t in tab.values() if t["source"] == "MISSING DATA")
    return {"figures": fig, "tables": tab, "stats": stats,
           "n_missing_figures": n_missing_fig, "n_missing_tables": n_missing_tab,
           "reports_dir": reports_dir}
